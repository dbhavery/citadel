"""File parsers for converting various document formats to plain text."""

import os
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""

    text: str
    metadata: dict = field(default_factory=dict)  # filename, extension, size, modified_at
    source_path: str = ""


# Extensions mapped to parser method names
_EXTENSION_MAP: dict[str, str] = {
    ".md": "_parse_markdown",
    ".txt": "_parse_text",
    ".py": "_parse_python",
    ".js": "_parse_code",
    ".ts": "_parse_code",
    ".go": "_parse_code",
    ".rs": "_parse_code",
    ".pdf": "_parse_pdf",
    ".docx": "_parse_docx",
}

# Extensions that are source code
_CODE_EXTENSIONS: dict[str, str] = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".go": "go",
    ".rs": "rust",
}


class DocumentParser:
    """Parse different file formats into plain text."""

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a file into a ParsedDocument based on its extension.

        Args:
            file_path: Absolute path to the file to parse.

        Returns:
            ParsedDocument with extracted text and metadata.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file extension is not supported.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        stat = path.stat()

        base_metadata = {
            "filename": path.name,
            "extension": ext,
            "size": stat.st_size,
            "modified_at": stat.st_mtime,
        }

        method_name = _EXTENSION_MAP.get(ext)
        if method_name is None:
            raise ValueError(f"Unsupported file extension: {ext}")

        method = getattr(self, method_name)
        text = method(path)

        # Add language metadata for code files
        if ext in _CODE_EXTENSIONS:
            base_metadata["language"] = _CODE_EXTENSIONS[ext]

        return ParsedDocument(
            text=text,
            metadata=base_metadata,
            source_path=str(path.resolve()),
        )

    def _parse_markdown(self, path: Path) -> str:
        """Parse a Markdown file into plain text."""
        return path.read_text(encoding="utf-8")

    def _parse_text(self, path: Path) -> str:
        """Parse a plain text file."""
        return path.read_text(encoding="utf-8")

    def _parse_python(self, path: Path) -> str:
        """Parse a Python source file, preserving content as-is."""
        return path.read_text(encoding="utf-8")

    def _parse_code(self, path: Path) -> str:
        """Parse a generic code file, preserving content as-is."""
        return path.read_text(encoding="utf-8")

    def _parse_pdf(self, path: Path) -> str:
        """Parse a PDF file into plain text.

        Requires PyPDF2 (optional dependency).
        """
        try:
            from PyPDF2 import PdfReader  # type: ignore[import-untyped] -- optional dep, may not have stubs
        except ImportError as exc:
            raise ImportError(
                "PyPDF2 is required to parse PDF files. "
                "Install it with: pip install citadel-ingest[pdf]"
            ) from exc

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def _parse_docx(self, path: Path) -> str:
        """Parse a DOCX file into plain text.

        Requires python-docx (optional dependency).
        """
        try:
            import docx  # type: ignore[import-untyped] -- optional dep, may not have stubs
        except ImportError as exc:
            raise ImportError(
                "python-docx is required to parse DOCX files. "
                "Install it with: pip install citadel-ingest[docx]"
            ) from exc

        doc = docx.Document(str(path))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return "\n\n".join(paragraphs)
